from __future__ import annotations

import tempfile
import unittest
from pathlib import Path
from unittest.mock import patch

from docx import Document

from scripts.docx import analyze_sample_pdf
from scripts.docx.generate_diagram_appendix_docx import extract_blocks
from scripts.docx.generate_thesis_docx import build_doc, load_style_profile
from scripts.workspace.check_thesis_workspace import check_optional
from scripts.workspace.init_workflow_logs import WORKFLOW_FILES, write_workflow_logs


class CoreContractTests(unittest.TestCase):
    def test_docx_generator_uses_builtin_styles_and_preserves_formula_text(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            source = Path(tmp) / "thesis.md"
            source.write_text(
                "# \u8bba\u6587\u9898\u76ee\n\n$$E = mc^2$$\n\n\u6b63\u6587\u7b2c\u4e00\u6bb5\u3002",
                encoding="utf-8",
            )

            document = build_doc(source, {}, load_style_profile())
            text = "\n".join(paragraph.text for paragraph in document.paragraphs)

            self.assertIn("\u8bba\u6587\u9898\u76ee", text)
            self.assertIn("E = mc^2", text)
            self.assertEqual(document.paragraphs[0].runs[0].font.size.pt, 18)
            self.assertEqual(document.sections[0].header.paragraphs[0].text, "")

    def test_docx_analyzer_outputs_lightweight_analysis_contract(self) -> None:
        from scripts.docx import analyze_docx

        with tempfile.TemporaryDirectory() as tmp:
            docx_path = Path(tmp) / "sample.docx"
            document = Document()
            document.sections[0].header.paragraphs[0].text = "School Header"
            document.add_heading("1.1 \u7814\u7a76\u80cc\u666f", level=2)
            document.add_paragraph("\u6b63\u6587\u5185\u5bb9")
            document.save(docx_path)

            payload = analyze_docx.analyze_docx(docx_path)

            self.assertEqual(payload["schema_version"], "1.0")
            self.assertEqual(payload["source"]["parser"], "analyze_docx.py")
            self.assertIn("outline_structure", payload["sample_patterns"])
            self.assertIn("style_observations", payload["sample_patterns"])
            self.assertTrue(
                any(
                    item["observation"] == "header_footer_text"
                    for item in payload["sample_patterns"]["style_observations"]
                )
            )

    def test_pdf_analyzer_falls_back_to_pdfplumber_text_structure(self) -> None:
        class FakeReader:
            outline = []

            def __init__(self, _path: str) -> None:
                self.pages = []

        class FakePlumberPage:
            def extract_text(self, **_kwargs) -> str:
                return "1.1 Research Background\nFigure 1-1 Architecture"

            def extract_tables(self) -> list[list[list[str]]]:
                return []

        class FakePdf:
            pages = [FakePlumberPage()]

            def __enter__(self):
                return self

            def __exit__(self, *_args) -> None:
                return None

        with tempfile.TemporaryDirectory() as tmp:
            pdf_path = Path(tmp) / "sample.pdf"
            pdf_path.write_bytes(b"%PDF-1.4\n")

            with patch.object(analyze_sample_pdf, "PdfReader", FakeReader):
                with patch.object(analyze_sample_pdf.pdfplumber, "open", return_value=FakePdf()):
                    payload = analyze_sample_pdf.analyze_pdf(pdf_path)

            self.assertEqual(payload["source"]["parser"], "analyze_sample_pdf.py")
            self.assertEqual(payload["confidence"]["structure"], "medium")
            self.assertEqual(payload["sample_patterns"]["outline_structure"][0]["title"], "1.1 Research Background")

    def test_workflow_initialization_and_optional_templates_are_lightweight(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            target = Path(tmp)

            written = write_workflow_logs(target)
            workflow = target / "paper-context" / "workflow"
            optional_result = check_optional(target, "templates/citation-crossref-register.yaml")

            self.assertEqual(set(written), {workflow / name for name in WORKFLOW_FILES})
            self.assertEqual(optional_result.status, "ok")

    def test_diagram_appendix_extracts_raw_plantuml_with_caption(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            source = Path(tmp) / "thesis.md"
            source.write_text(
                "@startuml\nAlice -> Bob: hello\n@enduml\n\n\u56fe3-1 \u767b\u5f55\u6d41\u7a0b\n",
                encoding="utf-8",
            )

            blocks = extract_blocks(source)

            self.assertEqual(blocks[0]["lang"], "plantuml")
            self.assertEqual(blocks[0]["caption"], "\u56fe3-1 \u767b\u5f55\u6d41\u7a0b")


if __name__ == "__main__":
    unittest.main()
